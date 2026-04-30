"""
文档解析模块 - 轻量级多格式文档解析
支持: PDF, Markdown, Word, CSV, Excel, TXT, JSON, HTML, EPUB
"""
from __future__ import annotations

import json
import mimetypes

from bs4 import BeautifulSoup
from chonkie import MarkdownChef, TableChef, TextChef
from docx import Document as DocxDocument
from ebooklib import epub
from pypdf import PdfReader


def _parse_pdf(file_path: str) -> list[dict]:
    """解析 PDF，返回每页文本"""
    reader = PdfReader(file_path)
    elements = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            elements.append({"type": "Page", "text": text.strip(), "metadata": {"page": i + 1}})
    return elements


def _parse_markdown(file_path: str) -> list[dict]:
    """使用 MarkdownChef 解析 md 文件，提取文本块、表格、代码块作为独立元素"""
    chef = MarkdownChef()
    doc = chef.process(file_path)
    elements = []

    for chunk in doc.chunks:
        elements.append({"type": "Text", "text": chunk.text, "metadata": {}})

    for table in doc.tables:
        elements.append({"type": "Table", "text": table.content, "metadata": {}})

    for code in doc.code:
        elements.append({"type": "Code", "text": code.content, "metadata": {"language": code.language}})

    return elements


def _parse_docx(file_path: str) -> list[dict]:
    """解析 Word 文档"""
    doc = DocxDocument(file_path)
    elements = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            elements.append({"type": "Paragraph", "text": text, "metadata": {"index": i}})
    for i, table in enumerate(doc.tables):
        table_text = ""
        for row in table.rows:
            table_text += " | ".join(cell.text.strip() for cell in row.cells) + "\n"
        if table_text.strip():
            elements.append({"type": "Table", "text": table_text.strip(), "metadata": {"index": i}})
    return elements


def _parse_csv(file_path: str) -> list[dict]:
    """使用 TableChef 解析 CSV，转为 markdown 表格格式"""
    chef = TableChef()
    doc = chef.process(file_path)
    elements = []
    for table in doc.tables:
        elements.append({"type": "Table", "text": table.content, "metadata": {}})
    return elements


def _parse_json(file_path: str) -> list[dict]:
    """解析 JSON 文件"""
    with open(file_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    return [{"type": "JSON", "text": json.dumps(data, ensure_ascii=False, indent=2), "metadata": {}}]


def _parse_html(file_path: str) -> list[dict]:
    """解析 HTML 文件，提取标题和段落"""
    with open(file_path, "r", encoding="utf-8") as f:
        soup = BeautifulSoup(f, "lxml")
    elements = []
    for tag in soup.find_all(["h1", "h2", "h3", "h4", "p", "li", "blockquote", "pre"]):
        text_content = tag.get_text(separator=" ", strip=True)
        if text_content:
            elements.append({"type": tag.name.upper(), "text": text_content, "metadata": {}})
    return elements


def _parse_epub(file_path: str) -> list[dict]:
    """解析 EPUB 文件"""
    book = epub.read_epub(file_path)
    elements = []
    for item in book.get_items():
        if item.get_type() == 1:
            soup = BeautifulSoup(item.get_content(), "html.parser")
            for tag in soup.find_all(["h1", "h2", "h3", "h4", "p", "li"]):
                text_content = tag.get_text(separator=" ", strip=True)
                if text_content:
                    elements.append({"type": tag.name.upper(), "text": text_content, "metadata": {}})
    return elements


def _parse_txt(file_path: str) -> list[dict]:
    """使用 TextChef 解析纯文本文件，返回全文作为单个 Text 元素（分块由 TextChunker 负责）"""
    chef = TextChef()
    doc = chef.process(file_path)
    return [{"type": "Text", "text": doc.content, "metadata": {}}]


_PARSERS = {
    "application/pdf": _parse_pdf,
    "text/markdown": _parse_markdown,
    "text/x-markdown": _parse_markdown,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": _parse_docx,
    "text/csv": _parse_csv,
    "application/vnd.ms-excel": _parse_csv,
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": _parse_csv,
    "text/plain": _parse_txt,
    "application/json": _parse_json,
    "text/html": _parse_html,
    "application/epub+zip": _parse_epub,
}


class DocumentParser:
    """轻量级多格式文档解析器"""

    def parse(self, file_path: str) -> list[dict]:
        """解析文档，返回结构化元素列表（含 Text/Table/Code/Paragraph 等类型）"""
        mime = mimetypes.guess_type(file_path)[0]
        parser = _PARSERS.get(mime)
        if parser:
            return parser(file_path)
        return [{"type": "NarrativeText", "text": open(file_path, encoding="utf-8").read().strip(), "metadata": {}}]

    def parse_text(self, text: str) -> list[dict]:
        """将纯文本转为元素列表（不分块，分块由 TextChunker 负责）"""
        return [{"type": "Text", "text": text, "metadata": {}}]
